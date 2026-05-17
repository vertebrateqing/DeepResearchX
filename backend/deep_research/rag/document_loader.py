from __future__ import annotations
"""Document loading for RAG ingest.

Supports PDF (pdfplumber → pypdf fallback), Word (.docx via python-docx),
and plain text (.txt / .md). Returns ``Document`` objects holding raw text
plus metadata. The actual splitting / embedding lives in
``deep_research.rag.pipeline``.
"""

import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Optional

logger = logging.getLogger(__name__)


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md", ".markdown"}


@dataclass
class Document:
    """A loaded document with raw text and metadata."""

    content: str
    metadata: dict[str, Any] = field(default_factory=dict)
    source: Optional[str] = None

    def to_dict(self) -> dict[str, Any]:
        return {
            "content": self.content,
            "metadata": dict(self.metadata),
            "source": self.source,
        }


# ---------------------------------------------------------------------------
# Format-specific loaders
# ---------------------------------------------------------------------------


def _load_pdf(file_path: Path) -> str:
    """Extract text from a PDF, preferring pdfplumber and falling back to pypdf."""
    try:
        import pdfplumber  # type: ignore

        texts: list[str] = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                if page_text.strip():
                    texts.append(page_text.strip())
        if texts:
            return "\n\n".join(texts)
    except ImportError:
        logger.debug("pdfplumber not installed, falling back to pypdf")
    except Exception as e:  # noqa: BLE001
        logger.warning(f"pdfplumber failed for {file_path}: {e}")

    try:
        from pypdf import PdfReader  # type: ignore

        reader = PdfReader(str(file_path))
        texts: list[str] = []
        for page in reader.pages:
            page_text = page.extract_text() or ""
            if page_text.strip():
                texts.append(page_text.strip())
        return "\n\n".join(texts)
    except ImportError as e:
        raise RuntimeError(
            "PDF support unavailable: install pdfplumber or pypdf"
        ) from e
    except Exception as e:  # noqa: BLE001
        logger.error(f"pypdf failed for {file_path}: {e}")
        return ""


def _load_docx(file_path: Path) -> str:
    """Extract text from a .docx file via python-docx."""
    try:
        from docx import Document as DocxDocument  # type: ignore
    except ImportError as e:
        raise RuntimeError(
            "DOCX support unavailable: install python-docx (pip install python-docx)"
        ) from e

    doc = DocxDocument(str(file_path))

    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Tables → tab-separated lines
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append("\t".join(cells))

    return "\n".join(parts)


def _load_text(file_path: Path) -> str:
    """Read a UTF-8 text file (.txt / .md)."""
    try:
        return file_path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        logger.warning(f"UTF-8 decode failed for {file_path}, trying gbk")
        return file_path.read_text(encoding="gbk", errors="ignore")


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def is_supported_file(file_name: str) -> bool:
    """Return True if the file extension is supported by load_document."""
    return Path(file_name).suffix.lower() in SUPPORTED_EXTENSIONS


def load_document(
    file_path: str | Path,
    extra_metadata: Optional[dict[str, Any]] = None,
) -> Document:
    """Load a document from disk and return a ``Document``.

    Args:
        file_path: path to a .pdf, .docx, .txt or .md file.
        extra_metadata: caller-provided metadata to merge in (e.g. doc_id,
            uploaded_at, original filename).

    Returns:
        Document with full text in ``content`` and metadata that includes
        ``filename``, ``extension``, ``size_bytes``, ``char_count``.

    Raises:
        FileNotFoundError: if file_path does not exist.
        ValueError: if extension is unsupported.
    """
    path = Path(file_path)
    if not path.exists() or not path.is_file():
        raise FileNotFoundError(f"Document not found: {path}")

    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type {ext!r}; supported: {sorted(SUPPORTED_EXTENSIONS)}"
        )

    if ext == ".pdf":
        content = _load_pdf(path)
    elif ext in {".docx", ".doc"}:
        content = _load_docx(path)
    else:
        content = _load_text(path)

    metadata: dict[str, Any] = {
        "filename": path.name,
        "extension": ext,
        "size_bytes": path.stat().st_size,
        "char_count": len(content),
    }
    if extra_metadata:
        metadata.update(extra_metadata)

    return Document(content=content, metadata=metadata, source=str(path))
